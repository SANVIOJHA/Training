from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"D:\Final_Sprint")
OUT = ROOT / "deliverables"
ASSETS = OUT / "assets"
LOGO = ASSETS / "fp_extracted" / "p1_0_Image1.png"
TRAINING_CERTIFICATE_IMAGE = Path(
    r"C:\Users\ASUS\OneDrive\Pictures\ScreenShotss\Screenshot 2026-06-03 223609.png"
)
FINAL = OUT / "SmartCourier_LPU_Internship_Report_Submission_Ready.docx"

BLUE = RGBColor(31, 78, 121)
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"


def set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def paragraph_border_bottom(paragraph, color="808080", size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def setup_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = BLUE
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10 if style_name == "Heading 1" else 6)
        style.paragraph_format.space_after = Pt(6)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_centered(doc, text, size=12, bold=False, after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p


def add_para(doc, text, after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_font(run, 12)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_font(run, 12)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_font(run, 12)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, bold=False, align=None, fill=None, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align or (WD_ALIGN_PARAGRAPH.CENTER if len(str(text)) < 24 else WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    set_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        shade_cell(cell, fill)


def add_table(doc, caption, headers, rows, widths=None):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(4)
    r = cap.add_run(caption)
    set_font(r, size=11, bold=True, italic=True)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, bold=True, fill=LIGHT_BLUE, size=10)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value, size=10)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_figure(doc, image_path: Path, caption: str, width=6.2):
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image_path), width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        set_font(r, size=11, bold=True, italic=True)


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_cover(doc):
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(3.2))
    add_centered(doc, "LOVELY PROFESSIONAL UNIVERSITY", 15, True, 2)
    add_centered(doc, "PHAGWARA, PUNJAB", 12, True, 18)
    add_centered(doc, "INTERNSHIP TRAINING REPORT", 14, True, 8)
    title = add_centered(doc, "SMARTCOURIER DELIVERY MANAGEMENT SYSTEM", 18, True, 6)
    paragraph_border_bottom(title, "1F4E79", "8")
    add_centered(doc, "Java Full Stack Training Project", 12, False, 14)
    add_centered(doc, "Submitted in partial fulfillment of the requirements for the award of degree of", 11, False, 2)
    add_centered(doc, "B.Tech Computer Science Engineering", 13, True, 18)
    add_centered(doc, "Training Organization: Capgemini", 12, True, 4)
    add_centered(doc, "Training Duration: 18 December 2025 to 21 May 2026", 12, False, 18)

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("SUBMITTED BY", "SUBMITTED TO"),
        ("Name of Student: ____________________\nRegistration Number: ____________________", "Name of Supervisor: ____________________\nDesignation: ____________________"),
        ("Signature of Student: ____________________", "Signature: ____________________"),
    ]
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            set_cell(table.rows[ri].cells[ci], val, bold=(ri == 0), fill=LIGHT_GRAY if ri == 0 else None, size=10)
    page_break(doc)


def front_matter(doc):
    doc.add_heading("Student Declaration", level=1)
    add_para(doc, "I hereby declare that this internship report titled SmartCourier Delivery Management System is an authentic record of the project work carried out by me during the Java Full Stack Intern training program at Capgemini from 18 December 2025 to 21 May 2026. The work described in this report is based on guided learning, implementation practice, testing, debugging, source-code study, and documentation completed during the training period.")
    add_para(doc, "SmartCourier is an academic training project prepared for evaluation and learning purposes. It is not presented as a confidential production client system. External documentation and technical references used while preparing the report are listed in the references section.")
    add_para(doc, "Date: ____________________\nPlace: ____________________\n\nSignature of Student: ____________________", after=0)
    page_break(doc)

    doc.add_heading("Declaration by Supervisors", level=1)
    add_para(doc, "This is to certify that the internship report titled SmartCourier Delivery Management System has been prepared by the student as part of the Java Full Stack training program. The report reflects the student's understanding of full stack development, microservices architecture, API design, database connectivity, messaging, security, testing, and professional documentation.")
    add_para(doc, "The work is suitable for academic evaluation as a training project and should be read in that context. This certificate is subject to verification and signature by the authorized faculty or industry supervisor.")
    add_para(doc, "Name of External Supervisor: ____________________        Name of Internal Supervisor: ____________________\nDesignation: ____________________                         Designation: ____________________\nSignature: ____________________                           Signature: ____________________\nDate: ____________________                                Date: ____________________", after=0)
    page_break(doc)

    doc.add_heading("Training Certificate", level=1)
    if TRAINING_CERTIFICATE_IMAGE.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(str(TRAINING_CERTIFICATE_IMAGE), width=Inches(5.65))
    else:
        add_para(doc, "This page is reserved for attaching the official training certificate issued by Capgemini or the authorized training department. The certificate may include the student's name, registration number, training role, internship duration, technology track, supervisor details, and organization seal.")
        add_para(doc, "[Attach scanned certificate or official letter here before final physical submission.]", after=0)
    page_break(doc)

    doc.add_heading("Acknowledgement", level=1)
    add_para(doc, "I express my sincere gratitude to Lovely Professional University for providing the academic framework and opportunity to undertake internship training as part of the B.Tech Computer Science Engineering curriculum. I am thankful to Capgemini for providing a professional training environment where I could study full stack development, enterprise application structure, and project-style execution.")
    add_para(doc, "I also thank my mentors, trainers, faculty guides, peers, and team members for their support during the training period. Their feedback helped me connect React, Spring Boot, REST APIs, microservices, MySQL, RabbitMQ, JWT security, testing, and debugging into one coherent project. The SmartCourier project gave me a practical way to apply these concepts in a realistic courier-management domain.")
    page_break(doc)

    doc.add_heading("Table of Contents", level=1)
    toc = [
        "Chapter 1: Introduction of the Company",
        "Chapter 2: Introduction of the Project Undertaken",
        "Chapter 3: Microservices Architecture and System Design",
        "Chapter 4: Implementation Details",
        "Chapter 5: Work Done During Training",
        "Chapter 6: Testing and Quality Assurance",
        "Chapter 7: Conclusion and Future Scope",
        "References",
    ]
    add_numbered(doc, toc)

    doc.add_heading("List of Figures", level=1)
    add_table(doc, "Figure Index", ["Figure No.", "Title"], [
        ("Figure 1.1", "Company organization structure in training context"),
        ("Figure 3.1", "SmartCourier microservices architecture"),
        ("Figure 3.2", "SmartCourier data flow diagram"),
        ("Figure 3.3", "SmartCourier use case diagram"),
        ("Figure 3.4", "Delivery booking and tracking sequence"),
        ("Figure 4.1", "SmartCourier entity relationship diagram"),
    ], [1.2, 5.2])

    doc.add_heading("List of Tables", level=1)
    add_table(doc, "Table Index", ["Table No.", "Title"], [
        ("Table 1.1", "Capgemini service areas"),
        ("Table 2.1", "User roles and responsibilities"),
        ("Table 2.2", "Functional requirements"),
        ("Table 2.3", "Non-functional requirements"),
        ("Table 3.1", "System components and responsibilities"),
        ("Table 4.1", "Technology stack and rationale"),
        ("Table 4.2", "Core API endpoint summary"),
        ("Table 5.1", "Training activities and learning outcomes"),
        ("Table 5.2", "Weekly learning summary"),
        ("Table 6.1", "Sample test cases and results"),
        ("Table 7.1", "Future enhancements"),
    ], [1.2, 5.2])

    doc.add_heading("List of Abbreviations", level=1)
    add_table(doc, "Abbreviation Index", ["Abbreviation", "Full Form"], [
        ("API", "Application Programming Interface"),
        ("DTO", "Data Transfer Object"),
        ("DFD", "Data Flow Diagram"),
        ("ER", "Entity Relationship"),
        ("HTTP", "Hypertext Transfer Protocol"),
        ("JPA", "Java Persistence API"),
        ("JWT", "JSON Web Token"),
        ("LPU", "Lovely Professional University"),
        ("OJT", "On-the-Job Training"),
        ("REST", "Representational State Transfer"),
        ("SPA", "Single Page Application"),
        ("UI", "User Interface"),
    ], [1.5, 4.9])
    page_break(doc)


def chapter_one(doc):
    doc.add_heading("Chapter 1: Introduction of the Company", level=1)
    add_para(doc, "Capgemini is a global business and technology transformation partner that works with organizations on consulting, digital engineering, cloud, data, application development, cybersecurity, and managed service programs. During my internship training, I used Capgemini's learning environment to understand how professional software teams organize technical work, communicate design decisions, test APIs, and document systems for review.")
    add_para(doc, "The organization traces its origin to 1967, when Serge Kampf founded Sogeti in Grenoble, France. Over time, Capgemini expanded into a multinational technology group through consulting, technology services, engineering, outsourcing, and digital transformation capabilities. This background helped me understand why modern software projects must balance business requirements, maintainability, security, delivery discipline, and user experience.")
    doc.add_heading("1.1 Vision, Mission and Work Culture", level=2)
    add_para(doc, "The training program connected Capgemini's larger technology-service context with day-to-day engineering practices: writing maintainable code, documenting API contracts, testing changes, following security principles, and explaining design choices clearly. The work culture visible during training emphasized teamwork, regular feedback, process discipline, and continuous learning.")
    add_para(doc, "These values influenced the SmartCourier project. I treated it as more than a coding assignment: the system needed a defined domain, clear module responsibilities, understandable documentation, and verifiable behavior that an evaluator could trace from the user interface to the backend services.")
    doc.add_heading("1.2 Services and Departments", level=2)
    add_para(doc, "A technology services organization includes delivery teams, architecture groups, quality assurance teams, cloud and infrastructure teams, security teams, human resources, training departments, finance, sales, and client engagement functions. In my training context, the most relevant groups were training coordination, mentors, technical trainers, peer groups, and evaluation support.")
    add_table(doc, "Table 1.1: Capgemini Service Areas", ["Area", "Description"], [
        ("Consulting", "Business and technology advisory support for transformation programs."),
        ("Application Development", "Design, development, modernization, and maintenance of software systems."),
        ("Cloud and Data", "Cloud migration, data platforms, analytics, AI-enabled solutions, and managed services."),
        ("Engineering", "Product engineering, intelligent industry, digital manufacturing, and platform engineering."),
        ("Cybersecurity", "Security assessment, identity management, threat monitoring, and secure architecture practices."),
    ], [1.8, 4.6])
    doc.add_heading("1.3 Training Environment", level=2)
    add_para(doc, "My internship training ran from 18 December 2025 to 21 May 2026. The work involved guided learning, hands-on implementation, peer discussion, mentor feedback, debugging, API testing, and project documentation. I worked primarily in the Java Full Stack track.")
    add_para(doc, "The tools and technologies practiced during the project included React, Vite, React Router, Context API, Axios, Spring Boot, Spring Cloud, MySQL, RabbitMQ, Eureka, OpenFeign, JWT, Postman, Zipkin, JUnit, Mockito, and MockMvc. Since SmartCourier was a training project, I present it as a realistic learning implementation rather than as a live client delivery assignment.")
    add_table(doc, "Figure 1.1: Company Organization Structure in Training Context", ["Training Area", "Role in Project Learning"], [
        ("Training Coordination", "Defined the training timeline, learning checkpoints, and submission expectations."),
        ("Technical Mentors", "Guided implementation concepts, debugging approach, and technology understanding."),
        ("Faculty/Evaluators", "Reviewed academic presentation, report quality, and project explanation."),
        ("Peer Group", "Supported discussion, testing practice, issue comparison, and viva preparation."),
    ], [2.2, 4.2])


def chapter_two(doc):
    doc.add_heading("Chapter 2: Introduction of the Project Undertaken", level=1)
    add_para(doc, "SmartCourier Delivery Management System is a web-based courier and parcel delivery training project. It models user registration, authentication, parcel booking, delivery assignment, shipment status updates, tracking history, hub management, administrative monitoring, and report generation. I built and studied the system to understand how a full stack application behaves when responsibilities are separated across a frontend, API Gateway, backend services, databases, and a message queue.")
    add_para(doc, "The original requirement document referred to an Angular frontend with Spring Boot microservices. The implemented source code uses React with Vite, React Router, Context API, Axios, and reusable components. I have documented this difference because the architectural intent remains the same: modular UI pages, protected routes, role-based access, REST API integration, and a service-oriented backend.")
    doc.add_heading("2.1 Problem Statement", level=2)
    add_para(doc, "Courier operations require accurate booking, pickup scheduling, shipment status updates, hub monitoring, exception handling, customer visibility, and administrative control. When these functions are handled manually or placed inside tightly coupled modules, errors can occur in status visibility, reporting, traceability, and access control. SmartCourier addresses this at training-project scale by separating authentication, delivery, tracking, and administration into focused services.")
    doc.add_heading("2.2 Objectives", level=2)
    add_bullets(doc, [
        "Digitize the parcel booking and delivery lifecycle for a courier-management workflow.",
        "Provide separate customer and administrator flows with role-based access.",
        "Implement REST APIs for authentication, delivery management, tracking, hub management, reports, and dashboard data.",
        "Use JWT authentication so protected APIs can identify users and enforce roles.",
        "Practice database-per-service design using MySQL and JPA/Hibernate.",
        "Use RabbitMQ events to decouple delivery status updates from tracking history persistence.",
        "Test APIs using Postman and backend tests, including success, validation, missing-token, unauthorized, and not-found scenarios.",
    ])
    doc.add_heading("2.3 Scope and Boundaries", level=2)
    add_para(doc, "The implemented scope covers customer signup and login, delivery booking, customer delivery listing, tracking lookup, delivery status updates, admin dashboard, user and hub management, report APIs, service discovery, API Gateway routing, RabbitMQ-based tracking updates, and API testing. The project also includes architecture diagrams, Postman collections, test-case documentation, viva preparation notes, and technical walkthrough material.")
    add_para(doc, "The project does not include production deployment, payment processing, live GPS devices, shipment label printing, courier partner API integration, or organization-specific client business rules. These limits were deliberate because the project was designed for academic evaluation and full stack training.")
    add_table(doc, "Table 2.1: User Roles and Responsibilities", ["Role", "Responsibilities"], [
        ("Customer", "Register, log in, book delivery, view personal deliveries, track shipment status, and access customer pages only after authentication."),
        ("Administrator", "View dashboard, manage deliveries, manage users, manage hubs, monitor reports, and update delivery status where required."),
        ("Delivery Agent (Extended Scope)", "Receive assigned deliveries and update shipment status in a future extension."),
        ("Trainer/Evaluator", "Review implementation, test APIs, inspect architecture, and evaluate project understanding."),
    ], [1.8, 4.6])
    add_table(doc, "Table 2.2: Functional Requirements", ["ID", "Requirement", "Module"], [
        ("FR-01", "User registration, login, refresh, validation, logout, and role lookup.", "Auth"),
        ("FR-02", "Create delivery with sender, receiver, source, destination, package weight, price, hub, and tracking details.", "Delivery"),
        ("FR-03", "Assign delivery to an agent, cancel, ship, deliver, and update status.", "Delivery"),
        ("FR-04", "Publish delivery status events and store tracking history.", "Delivery/Tracking"),
        ("FR-05", "Fetch deliveries by customer, id, tracking number, status, agent, price threshold, and search parameters.", "Delivery"),
        ("FR-06", "Manage logistics hubs with create, update, delete, and list operations.", "Delivery"),
        ("FR-07", "Create, update, list, filter, count, and summarize reports.", "Admin"),
        ("FR-08", "Show admin dashboard counts using service-to-service communication.", "Admin"),
    ], [0.8, 4.5, 1.1])
    add_table(doc, "Table 2.3: Non-Functional Requirements", ["Attribute", "Expected Behavior"], [
        ("Security", "JWT-protected APIs and role-aware frontend routes restrict access."),
        ("Maintainability", "Code is organized into controller, service, repository, DTO, entity, exception, and configuration layers."),
        ("Scalability", "Microservices can be developed and scaled independently in the architecture model."),
        ("Reliability", "RabbitMQ supports event-driven tracking updates and reduces direct coupling."),
        ("Usability", "Frontend workflows guide users through login, booking, tracking, dashboard, and management pages."),
        ("Observability", "Eureka and Zipkin support service visibility and request tracing during local evaluation."),
    ], [1.8, 4.6])


def chapter_three(doc):
    doc.add_heading("Chapter 3: Microservices Architecture and System Design", level=1)
    add_para(doc, "SmartCourier follows a microservices-based architecture. The React frontend communicates with the Spring Cloud API Gateway, and the gateway routes requests to individual services. Each service owns a specific domain and uses its own persistence boundary. Eureka provides service discovery, RabbitMQ supports asynchronous events, OpenFeign supports synchronous service-to-service calls, and Zipkin helps trace requests during debugging.")
    add_figure(doc, ASSETS / "architecture.png", "Figure 3.1: SmartCourier Microservices Architecture")
    add_table(doc, "Table 3.1: System Components and Responsibilities", ["Component", "Responsibility", "Learning Outcome"], [
        ("React Frontend", "SPA screens, forms, protected routes, dashboards, and Axios API calls.", "Component design and user flow."),
        ("API Gateway", "Single external entry point for /gateway paths, token validation, and routing.", "Centralized security and routing."),
        ("Eureka Server", "Registers services and allows lookup by logical service name.", "Dynamic service discovery."),
        ("Config Server", "Centralizes configuration for services during local setup.", "Configuration management."),
        ("Auth Service", "Registration, login, JWT creation, user APIs, role lookup, validation, and logout.", "Security and identity."),
        ("Delivery Service", "Booking, pricing, hub management, assignment, status changes, and delivery search.", "Core domain logic."),
        ("Tracking Service", "Tracking events, latest status, range queries, counts, and event consumption.", "Event-driven state history."),
        ("Admin Service", "Reports, dashboard aggregation, delivery monitoring, and admin status updates.", "Service aggregation and reports."),
        ("RabbitMQ", "Message exchange and queues for delivery and user events.", "Asynchronous communication."),
        ("MySQL", "Relational storage for each service schema.", "Persistence and JPA mapping."),
        ("Zipkin", "Distributed tracing across gateway and services.", "Debugging request paths."),
    ], [1.55, 3.1, 1.55])
    doc.add_heading("3.1 Request Flow", level=2)
    add_numbered(doc, [
        "The user interacts with the React frontend by logging in, booking a parcel, viewing deliveries, tracking a shipment, or opening an admin page.",
        "Axios sends the request to the gateway base URL, http://localhost:8080/gateway.",
        "The gateway validates JWT information for protected routes and forwards the request to the target service.",
        "Eureka allows the gateway and Feign clients to locate services by logical service name instead of hardcoded service addresses.",
        "The target service executes business logic through controller, service, repository, entity, DTO, mapper, and exception-handling layers.",
        "When a delivery is created or its status changes, the Delivery Service publishes an event to RabbitMQ.",
        "The Tracking Service consumes delivery status events and stores tracking history in its own database.",
        "The Admin Service can fetch delivery data synchronously through OpenFeign to build dashboard summaries.",
    ])
    doc.add_heading("3.2 Data Flow and Use Cases", level=2)
    add_para(doc, "The data flow separates user input, authentication, delivery processing, tracking updates, and report generation. Customer use cases focus on signup, login, booking, viewing personal deliveries, and tracking. Administrator use cases focus on delivery monitoring, user and hub management, report generation, and dashboard visibility.")
    add_figure(doc, ASSETS / "dfd.png", "Figure 3.2: SmartCourier Data Flow Diagram")
    add_figure(doc, ASSETS / "use_case.png", "Figure 3.3: SmartCourier Use Case Diagram", width=6.0)
    doc.add_heading("3.3 Synchronous and Asynchronous Communication", level=2)
    add_para(doc, "Synchronous communication is used when the caller needs an immediate response. The frontend waits for login results, delivery details, tracking results, and dashboard data. The Admin Service also uses OpenFeign to fetch delivery information from the Delivery Service, while Eureka resolves the actual service instance.")
    add_para(doc, "Asynchronous communication is used when services should not block each other. The Delivery Service publishes events through RabbitMQ after booking creation or status updates. The Tracking Service listens to the configured queue and stores tracking records separately. This design keeps tracking history available without making the Delivery Service directly responsible for tracking persistence.")
    add_figure(doc, ASSETS / "sequence.png", "Figure 3.4: Delivery Booking and Tracking Sequence")


def chapter_four(doc):
    doc.add_heading("Chapter 4: Implementation Details", level=1)
    doc.add_heading("4.1 Technology Choices", level=2)
    add_table(doc, "Table 4.1: Technology Stack and Rationale", ["Layer", "Technology", "Reason for Selection"], [
        ("Frontend", "React, Vite, React Router, Context API, Axios", "Fast SPA development, component-based UI, protected routing, global auth state, and HTTP integration."),
        ("Styling/UI", "CSS/Tailwind-style utility conventions and reusable components", "Consistent layout for landing, auth, customer, admin, and tracking pages."),
        ("Backend", "Java 17+, Spring Boot 3.x, Spring Cloud", "REST API development with gateway, discovery, Feign, security, and messaging support."),
        ("Security", "Spring Security, JWT, BCrypt", "Stateless authentication, role-aware access, and password hashing."),
        ("Persistence", "MySQL, JPA/Hibernate, Flyway", "Relational modeling, repository pattern, transactions, entity mapping, and repeatable migrations."),
        ("Messaging", "RabbitMQ", "Decouples delivery status changes from tracking persistence."),
        ("Service Discovery", "Netflix Eureka", "Allows services and gateway to use logical service names."),
        ("Observability", "Zipkin/Micrometer tracing", "Helps inspect request flow across multiple services."),
        ("Testing", "JUnit, Mockito, MockMvc, Postman", "Supports unit, controller, integration, and manual API validation."),
    ], [1.25, 2.35, 2.8])
    doc.add_heading("4.2 Frontend Implementation", level=2)
    add_para(doc, "The frontend is implemented as a React single page application. Routing is centralized in AppRoutes.jsx. Public routes include the landing page, about page, login, and signup. Customer routes such as /customer/dashboard, /customer/deliveries/new, and /customer/track/:trackingNumber are protected with the CUSTOMER role. Admin routes such as /admin/dashboard, /admin/deliveries, /admin/users, and /admin/hubs are protected with the ADMIN role.")
    add_para(doc, "Authentication state is handled through AuthContext. The context keeps the logged-in user state and provides login, logout, and authentication-check behavior to navigation, public routes, and protected routes. This avoids prop drilling and keeps security state available wherever it is required.")
    add_para(doc, "API integration is handled through an Axios instance configured with the gateway base URL. The request interceptor attaches the JWT token to outgoing requests, while the response logic can redirect the user when authorization fails. This keeps token handling centralized instead of repeating header logic inside every page.")
    doc.add_heading("4.3 Backend Implementation", level=2)
    add_para(doc, "The backend is divided into services. Each service follows a layered Spring Boot structure: controller classes expose REST endpoints, service classes contain business logic, repository interfaces handle database access, entities map to database tables, DTOs shape request and response payloads, and exception handlers return consistent error responses.")
    add_para(doc, "The Auth Service implements signup, login, refresh, token validation, current-user lookup, role lookup, user listing, user update, role update, deletion, and logout. The User entity stores username, email, phone number, encrypted password, role, active status, created timestamp, and updated timestamp. BCrypt is used so raw passwords are not stored.")
    add_para(doc, "The Delivery Service is the main business service. It stores delivery records with tracking number, customer username, sender and receiver details, source, destination, status, assigned agent, current hub, price, weight, and timestamps. It also manages hubs with name, code, city, state, address, contact number, and audit timestamps. Pricing is implemented through a Strategy Pattern in which PricingStrategyFactory selects standard or express pricing logic based on delivery characteristics.")
    add_para(doc, "The Tracking Service stores tracking events with delivery id, status, location, description, and timestamp. It exposes APIs for creating tracking events, retrieving all events for a delivery, latest status, status/location filters, date-range filtering, counts, existence checks, deletion, and health checks. It also consumes RabbitMQ messages so delivery status changes become tracking history.")
    add_para(doc, "The Admin Service manages reports and dashboard-related operations. Reports contain report type, report date, from/to range, report data, and generated-by information. The service exposes report CRUD, filtering, count, latest report, type summaries, dashboard, delivery monitoring, and admin delivery status update APIs.")
    doc.add_heading("4.4 Database Design", level=2)
    add_para(doc, "SmartCourier follows database-per-service ownership. Auth data is owned by the Auth Service, delivery and hub data by the Delivery Service, tracking history by the Tracking Service, and report data by the Admin Service. This prevents one service from directly modifying another service's tables and keeps domain boundaries clear.")
    add_figure(doc, ASSETS / "er.png", "Figure 4.1: SmartCourier Entity Relationship Diagram")
    doc.add_heading("4.5 API Design", level=2)
    add_para(doc, "The APIs are grouped by service and exposed through the gateway under /gateway. Direct service URLs are also useful during debugging: Auth runs on port 8081, Delivery on 8082, Tracking on 8083, Admin on 8084, and the gateway on 8080. The test documentation identifies 62 implemented endpoints across the four business services.")
    add_table(doc, "Table 4.2: Core API Endpoint Summary", ["Service", "Endpoint Group", "Purpose"], [
        ("Auth", "POST /auth/signup, POST /auth/login, POST /auth/refresh", "Registration, login, and token refresh."),
        ("Auth", "GET /auth/validate, /role, /me, /users/*", "Token validation, current user, role lookup, and user management."),
        ("Delivery", "GET /deliveries/my, POST /deliveries, GET /deliveries/{id}", "Customer delivery listing, booking, and detail retrieval."),
        ("Delivery", "PUT /deliveries/{id}/status, /assign, /cancel, /ship, /deliver", "Lifecycle and assignment operations."),
        ("Delivery", "GET /deliveries/tracking/{trackingNumber}, /summary/status, /search", "Tracking lookup, status summaries, and filtering."),
        ("Delivery", "/deliveries/hubs", "Hub create, update, delete, detail, and list operations."),
        ("Tracking", "POST /tracking, GET /tracking/{deliveryId}, /latest, /range", "Tracking event creation and retrieval."),
        ("Tracking", "/status/{status}, /location/{location}, /count, /exists", "Tracking filters and status checks."),
        ("Admin", "/admin/reports, /admin/reports/latest, /summary/type", "Report management and summaries."),
        ("Admin", "/admin/dashboard, /admin/deliveries, /admin/deliveries/{id}/status", "Dashboard, delivery monitoring, and admin status updates."),
    ], [1.1, 3.1, 2.2])
    doc.add_heading("4.6 Security Implementation", level=2)
    add_para(doc, "Security is implemented through JWT-based authentication. After a successful login, the Auth Service returns a token. The frontend stores the token and sends it in the Authorization: Bearer <token> header for protected API calls. The API Gateway validates the token before forwarding secured requests. Role-based logic is applied in protected frontend routes and backend access checks.")
    add_para(doc, "For downstream services, the gateway can propagate identity details through headers such as username or role. The Delivery Service uses customer identity to return only the current user's deliveries through /deliveries/my and related endpoints. This avoids relying on user-submitted customer identifiers for protected customer data.")


def chapter_five(doc):
    doc.add_heading("Chapter 5: Work Done During Training", level=1)
    add_para(doc, "During the internship, I worked through SmartCourier as an end-to-end learning exercise. I studied the requirement document, understood the courier domain, set up the frontend and backend, connected services through gateway routes, practiced database persistence, implemented and tested APIs, and prepared documentation for evaluation.")
    add_table(doc, "Table 5.1: Training Activities and Learning Outcomes", ["Activity", "Work Performed", "Learning Outcome"], [
        ("Orientation", "Understood training objectives, tools, project expectations, and the full stack roadmap.", "Prepared environment and learning notes."),
        ("React Practice", "Worked on pages, forms, routing, protected routes, dashboards, and reusable components.", "Improved UI component and workflow understanding."),
        ("Spring Boot Practice", "Created controllers, services, repositories, DTOs, entities, exception handlers, and configuration classes.", "Understood layered backend design."),
        ("Database Connectivity", "Configured MySQL, JPA entities, repositories, and service-level persistence.", "Practiced schema design and repository pattern."),
        ("Authentication", "Implemented and tested signup, login, JWT token usage, protected APIs, and role checks.", "Understood stateless security workflow."),
        ("RabbitMQ", "Published delivery status events and consumed tracking updates.", "Learned asynchronous service communication."),
        ("Testing", "Used Postman, unit tests, controller tests, validation checks, and integration checklists.", "Improved debugging and QA discipline."),
        ("Documentation", "Prepared report, architecture notes, viva question bank, walkthrough, and test-case documentation.", "Improved technical communication."),
    ], [1.45, 3.05, 1.9])
    doc.add_heading("5.1 Delivery Service Deep Dive", level=2)
    add_para(doc, "The Delivery Service was the most important module because it holds the core courier workflow. A booking request enters through the controller, is validated through DTO constraints, passes into the service layer, calculates pricing through the pricing strategy, saves the record through JPA, and publishes an event for tracking. This module helped me understand how business logic, validation, persistence, and messaging fit together.")
    add_para(doc, "Important decisions in this module include using an enum for delivery status instead of a free-form string, generating a unique tracking number, storing the authenticated customer username, linking the current hub through a JPA relationship, and using @Transactional service methods so related changes either complete together or fail safely.")
    doc.add_heading("5.2 Challenges Faced and Solutions Implemented", level=2)
    add_bullets(doc, [
        "Challenge: Understanding request flow across frontend, gateway, services, database, and queue. Solution: traced one workflow at a time from UI action to API, service method, database record, RabbitMQ message, and tracking result.",
        "Challenge: Handling the React-versus-Angular difference between the specification and implementation. Solution: preserved the original requirement intent while documenting the actual React implementation accurately.",
        "Challenge: Debugging protected requests. Solution: checked token creation, browser storage, Axios headers, gateway validation, role checks, and downstream service headers in sequence.",
        "Challenge: Avoiding tightly coupled services. Solution: used RabbitMQ for status events and OpenFeign only where immediate dashboard data was required.",
        "Challenge: Validating APIs beyond happy paths. Solution: created Postman and test-case coverage for missing tokens, invalid bodies, invalid ids, duplicate users, unauthorized roles, and unavailable downstream services.",
        "Challenge: Explaining microservices clearly for viva. Solution: prepared architecture diagrams, walkthrough notes, endpoint summaries, and module-level explanations.",
    ])
    doc.add_heading("5.3 Weekly Learning Summary", level=2)
    add_table(doc, "Table 5.2: Weekly Learning Summary", ["Week", "Learning Focus", "Work Summary"], [
        ("Week 1", "Java refresh and OOP", "Practiced core Java concepts and prepared for backend work."),
        ("Week 2", "Spring Boot basics", "Created basic REST controllers and understood application structure."),
        ("Week 3", "REST API practice", "Worked with HTTP methods, JSON payloads, status codes, and Postman."),
        ("Week 4", "React components", "Built and reviewed reusable UI components."),
        ("Week 5", "Routing and Context API", "Implemented protected routes and global authentication state."),
        ("Week 6", "MySQL and JPA", "Mapped entities, repositories, and database operations."),
        ("Week 7", "JWT security", "Studied login, token generation, token validation, and role checks."),
        ("Week 8", "Microservices and Eureka", "Registered services and understood discovery flow."),
        ("Week 9", "RabbitMQ messaging", "Published and consumed events for tracking updates."),
        ("Week 10", "OpenFeign and admin dashboard", "Practiced service-to-service communication and fallback thinking."),
        ("Week 11", "Postman testing", "Prepared API requests for integrated testing."),
        ("Week 12", "Debugging and exception handling", "Handled validation and not-found scenarios more consistently."),
        ("Week 13", "Frontend integration", "Connected UI flows with gateway APIs."),
        ("Week 14", "Unit testing", "Practiced JUnit, Mockito, and controller tests."),
        ("Week 15", "Documentation", "Created report sections and architecture notes."),
        ("Week 16", "Viva preparation", "Prepared question bank and walkthrough material."),
        ("Week 17", "Final review", "Reviewed source, APIs, diagrams, and testing evidence."),
        ("Week 18", "Future scope study", "Documented cloud, GPS, notifications, and route optimization improvements."),
        ("Week 19", "Report drafting", "Organized chapters, tables, figures, and references."),
        ("Week 20", "Submission readiness", "Reviewed formatting, grammar, duplication, and evaluator flow."),
    ], [1.0, 2.0, 3.4])


def chapter_six(doc):
    doc.add_heading("Chapter 6: Testing and Quality Assurance", level=1)
    add_para(doc, "Testing focused on verifying individual services, API contracts, secured request behavior, and cross-service workflows. Since SmartCourier is a training project, I used a practical combination of automated backend tests, controller tests, Postman collections, validation checks, and manual UI workflow verification.")
    doc.add_heading("6.1 Testing Strategy", level=2)
    add_para(doc, "Unit tests were used for service-layer behavior where dependencies could be mocked. Mockito helped isolate logic that depended on repositories, mappers, pricing strategies, RabbitMQ publishers, or Feign clients. Controller tests checked whether REST endpoints accepted input and returned expected HTTP responses.")
    add_para(doc, "Postman testing was used for integrated API verification. The collection uses the gateway base URL and variables for access token, refresh token, usernames, password, phone number, delivery id, tracking number, hub id, report id, and assigned agent. Login and refresh flows can save tokens for later protected requests.")
    add_para(doc, "Integration-focused checks covered Auth plus Gateway, Delivery plus RabbitMQ plus Tracking, Admin plus Feign plus Delivery, Eureka registration, Config Server access, and Zipkin traces. These checks helped me understand failures caused by configuration, service discovery, message queues, token headers, or database state.")
    add_table(doc, "Table 6.1: Sample Test Cases and Results", ["ID", "Scenario", "Expected Result", "Status"], [
        ("TC-01", "Signup with valid username, phone number, and password.", "User registered successfully.", "Pass"),
        ("TC-02", "Login with valid credentials.", "JWT access token returned.", "Pass"),
        ("TC-03", "Access protected delivery API without token.", "401 Unauthorized.", "Pass"),
        ("TC-04", "Create delivery with valid request.", "Delivery created with CREATED status and tracking number.", "Pass"),
        ("TC-05", "Update delivery status to SHIPPED.", "Status updated and RabbitMQ event published.", "Pass"),
        ("TC-06", "Consume delivery status event.", "Tracking event stored by Tracking Service.", "Pass"),
        ("TC-07", "Customer opens admin route.", "403 Forbidden or frontend redirect.", "Pass"),
        ("TC-08", "Admin dashboard fetches counts.", "Dashboard values returned or safe fallback shown.", "Pass"),
        ("TC-09", "Invalid delivery id lookup.", "404 Not Found or handled error response.", "Pass"),
        ("TC-10", "Eureka dashboard after startup.", "Services shown as UP.", "Pass"),
    ], [0.75, 2.45, 2.5, 0.7])
    doc.add_heading("6.2 Evaluation Demo Flow", level=2)
    add_numbered(doc, [
        "Show Eureka with gateway, config server, and business services running.",
        "Open Config Server URL to confirm service configuration is available.",
        "Sign up or log in through the Auth API.",
        "Use the returned JWT on gateway delivery APIs.",
        "Create a delivery and note the tracking number.",
        "Update delivery status and observe RabbitMQ queue activity.",
        "Fetch tracking history to confirm the event was stored.",
        "Open the admin dashboard to verify dashboard counts.",
        "Show Zipkin traces for a gateway request and downstream service call.",
        "Show Swagger or Postman requests for selected services.",
    ])
    doc.add_heading("6.3 Results Achieved", level=2)
    add_para(doc, "The final training implementation demonstrates a working full stack courier-management flow. Customers can authenticate, book deliveries, view their deliveries, and track status. Administrators can access protected management views and dashboard/report APIs. The backend demonstrates 62 service endpoints across Auth, Delivery, Tracking, and Admin services. RabbitMQ-based events update tracking records, while OpenFeign supports administrative aggregation.")
    add_para(doc, "The project also produced supporting artifacts: architecture diagrams, ER diagram, sequence diagram, data-flow diagram, use-case diagram, Postman collection, test-case documentation, viva question bank, technical concept explanations, and project walkthrough notes. These outputs made the project easier to evaluate and explain.")


def chapter_seven(doc):
    doc.add_heading("Chapter 7: Conclusion and Future Scope", level=1)
    add_para(doc, "The internship training helped me connect academic knowledge with practical full stack software development. Through SmartCourier, I understood how a courier-management system can be divided into clear responsibilities: the frontend handles user workflows, the gateway handles routing and token checks, each service owns its domain logic, repositories manage persistence, and messaging supports asynchronous updates.")
    add_para(doc, "The project improved my confidence in React, Spring Boot, REST APIs, JWT security, MySQL, JPA/Hibernate, RabbitMQ, Eureka, OpenFeign, Postman, Zipkin, and unit testing concepts. More importantly, it helped me explain request flow, justify architecture choices, test APIs systematically, and debug integration issues across multiple layers.")
    doc.add_heading("7.1 Key Learnings", level=2)
    add_bullets(doc, [
        "A full stack feature must be understood from the user interface through the gateway, service, database, and response.",
        "JWT security is not only a login feature; it affects routing, API headers, role checks, protected screens, and error handling.",
        "Microservices improve separation of responsibilities but introduce configuration, discovery, messaging, and debugging complexity.",
        "RabbitMQ is useful when one service needs to notify another without waiting for that service to complete work immediately.",
        "Testing must include success cases, invalid input, missing authorization, role restrictions, not-found records, and service integration scenarios.",
        "Good documentation makes technical work easier to evaluate, maintain, and explain in viva.",
    ])
    doc.add_heading("7.2 Future Scope", level=2)
    add_table(doc, "Table 7.1: Future Enhancements", ["Enhancement", "Description"], [
        ("AI-Based Route Optimization", "Recommend efficient routes using distance, delivery priority, hub capacity, traffic data, and historical patterns."),
        ("Real-Time GPS Tracking", "Show live parcel or delivery-agent location using map integration and privacy-aware location updates."),
        ("Payment Integration", "Support online payment, invoices, refunds, and secure transaction records."),
        ("Notification System", "Send email, SMS, or push notifications for booking confirmation, status changes, delays, and delivery proof."),
        ("Shipment Labels and Documents", "Generate labels, invoices, proof-of-delivery documents, and upload supporting files."),
        ("Cloud Deployment", "Containerize services and deploy with managed databases, broker, secrets management, logging, and monitoring."),
        ("CI/CD and Quality Gates", "Automate build, test, code quality checks, and deployment pipelines."),
        ("Centralized Logging and Metrics", "Add dashboards for service health, latency, error rates, queue depth, and trace analysis."),
        ("API Rate Limiting", "Protect public endpoints from abuse and improve gateway-level security."),
    ], [2.1, 4.3])
    doc.add_heading("7.3 Final Conclusion", level=2)
    add_para(doc, "SmartCourier remained within the scope of an academic and training project, but it successfully modeled realistic logistics workflows and modern application design patterns. The project demonstrates how authentication, delivery booking, tracking, administration, database design, microservices communication, security, testing, and documentation can be combined into one coherent system. It strengthened my ability to build, analyze, and explain enterprise-style full stack applications.")


def references(doc):
    doc.add_heading("References", level=1)
    refs = [
        "Capgemini. About us. https://www.capgemini.com/us-en/about-us/",
        "Capgemini. What we do: services and solutions. https://www.capgemini.com/us-en/about-us/who-we-are/what-we-do/",
        "Capgemini. Who we are. https://www.capgemini.com/gb-en/about-us/who-we-are/",
        "Spring Boot Reference Documentation. https://docs.spring.io/spring-boot/docs/current/reference/html/",
        "Spring Framework Documentation. https://docs.spring.io/spring-framework/reference/",
        "Spring Cloud Gateway Documentation. https://docs.spring.io/spring-cloud-gateway/reference/",
        "Spring Cloud Netflix Eureka Documentation. https://docs.spring.io/spring-cloud-netflix/reference/",
        "Spring Data JPA Documentation. https://docs.spring.io/spring-data/jpa/reference/",
        "React Documentation. https://react.dev/",
        "React Router Documentation. https://reactrouter.com/",
        "Axios Documentation. https://axios-http.com/docs/intro",
        "Oracle Java Documentation. https://docs.oracle.com/en/java/",
        "RabbitMQ Documentation. https://www.rabbitmq.com/docs",
        "MySQL Documentation. https://dev.mysql.com/doc/",
        "JSON Web Token Introduction. https://jwt.io/introduction",
        "OpenAPI Specification. https://spec.openapis.org/oas/latest.html",
        "Postman Learning Center. https://learning.postman.com/docs/",
        "Hibernate ORM Documentation. https://hibernate.org/orm/documentation/",
    ]
    add_numbered(doc, refs)


def main():
    doc = Document()
    setup_document(doc)
    add_cover(doc)
    front_matter(doc)
    chapter_one(doc)
    chapter_two(doc)
    chapter_three(doc)
    chapter_four(doc)
    chapter_five(doc)
    chapter_six(doc)
    chapter_seven(doc)
    references(doc)
    doc.save(FINAL)
    print(FINAL)


if __name__ == "__main__":
    main()
