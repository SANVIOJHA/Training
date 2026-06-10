from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(r"D:\Final_Sprint")
SRC = ROOT / "deliverables" / "SmartCourier_LPU_Internship_Report.docx"
OUT = ROOT / "deliverables" / "SmartCourier_LPU_Internship_Report_Revised.docx"


def replace_paragraph(paragraph, text):
    """Replace paragraph text while preserving its paragraph style."""
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def replace_heading(paragraph, text):
    replace_paragraph(paragraph, text)


def main():
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)

    replacements = {
        # Chapter 1
        41: "Capgemini is a global consulting, technology, and engineering services organization. For this internship report, I studied the company as the professional context in which my Java Full Stack training was conducted, rather than as a project client.",
        42: "The company originated in France in 1967 and grew through decades of work in IT services, consulting, outsourcing, cloud, engineering, and digital transformation. Its long presence in enterprise technology helped me understand why software projects are expected to follow disciplined delivery practices.",
        43: "During training, I connected the company profile with the way technical learning was organized: structured modules, practical assignments, review discussions, and an emphasis on reliable implementation. This made the company introduction relevant to my internship experience.",
        44: "Capgemini's service portfolio includes consulting, application development, cloud transformation, data and analytics, cybersecurity, digital engineering, enterprise platforms, and managed services. These areas show how a modern IT organization supports both business and technical change.",
        45: "The departments I observed indirectly through training included learning and development, delivery, engineering, quality assurance, infrastructure, security, and project coordination. Each function contributes to making software delivery systematic instead of depending only on individual coding effort.",
        46: "My training environment was based on guided learning and hands-on development. I worked with frontend, backend, database, security, messaging, and testing concepts through the SmartCourier project, which helped me relate company-level technology areas to a realistic trainee assignment.",
        49: "The organizational structure can be understood as a hierarchy of leadership, regional units, service lines, delivery teams, project teams, and support functions. For a trainee, this structure explains how learning, mentoring, project execution, and review activities are coordinated.",
        53: "Capgemini began as Sogeti, founded by Serge Kampf in Grenoble, France, in 1967. The organization gradually expanded beyond its original regional base and became known for technology services, consulting, and business transformation support.",
        54: "Its growth was influenced by major changes in enterprise computing: mainframe services, software development, outsourcing, cloud platforms, digital channels, data-driven decision-making, and engineering services. This history shows that the company evolved with the changing needs of clients and technology markets.",
        55: "I included this background because it gives context to the internship setting. A company with decades of technology experience is likely to emphasize process, documentation, teamwork, security, and maintainability, which were also reflected in my training work.",
        57: "Capgemini's vision is associated with helping organizations use technology to achieve practical business transformation. I understood this not as a slogan, but as the idea that software should solve operational problems, improve visibility, and support better decisions.",
        58: "The mission and values of the company are reflected in client focus, collaboration, responsible innovation, respect for people, and professional integrity. In my training, these values appeared in simple habits such as explaining assumptions clearly, accepting feedback, and improving the solution step by step.",
        59: "The values most relevant to my experience were teamwork, honesty in reporting progress, curiosity while learning unfamiliar tools, and accountability for assigned work. These values helped me approach SmartCourier as a learning project with professional seriousness.",
        61: "A large IT services company normally works through specialized departments. Consulting teams identify business requirements, delivery teams build and maintain solutions, engineering teams handle technical design, quality teams verify correctness, and support teams keep systems stable after release.",
        62: "Other functions such as cloud, cybersecurity, data, human resources, finance, and learning and development support the core delivery work. From my trainee perspective, the learning and mentoring functions were the most visible because they shaped the curriculum, assignments, and evaluation discussions.",
        63: "This section helped me understand that a software product is not produced by developers alone. Requirements, architecture, testing, deployment, documentation, and coordination all depend on different responsibilities working together.",
        65: "Capgemini works across several technology focus areas, including cloud computing, artificial intelligence, data analytics, cybersecurity, digital engineering, enterprise applications, automation, and modern software development.",
        66: "Cloud and DevOps practices support scalable deployment and environment management. AI and data analytics help organizations use information more intelligently. Cybersecurity protects identities, APIs, data, and infrastructure. Digital engineering connects software with product and operational systems.",
        67: "The SmartCourier project aligned most closely with application development, API design, authentication, database usage, microservices, service discovery, and messaging. These areas gave me a manageable way to practice technologies that are widely used in enterprise projects.",
        69: "During my internship at Capgemini from 18 December 2025 to 21 May 2026, I followed a structured Java Full Stack training path. The learning setup combined concept sessions, implementation practice, doubt clearing, testing activities, and documentation work.",
        70: "My training setup included React for the frontend, Spring Boot for backend APIs, MySQL for persistence, RabbitMQ for asynchronous communication, JWT for security, and supporting tools such as Postman, Git, and Maven. I used SmartCourier to connect these topics into one workflow.",
        71: "The environment encouraged gradual improvement. When an error occurred, I was expected to check logs, review configurations, test the API separately, and explain the issue clearly. This made the training experience more practical than simply reading theory.",

        # Chapter 2
        74: "SmartCourier Delivery Management System was the training project I used to apply Java Full Stack concepts in a logistics-style workflow. The project simulated how customers book deliveries, how administrators manage operational data, and how tracking information is updated.",
        75: "The problem addressed by the project is the difficulty of coordinating courier operations when booking, authentication, delivery status, tracking history, and administrative reporting are handled separately or manually.",
        76: "My objective was to build and understand an end-to-end academic implementation. On the frontend, I focused on screens, routing, forms, role-based access, and API calls. On the backend, I worked with REST endpoints, service layers, repositories, security, messaging, and databases.",
        77: "The project includes customer and administrator roles. Customers can register, log in, create delivery requests, view their own deliveries, and track shipment progress. Administrators can review deliveries, manage users or hubs in the training scope, and access dashboard-style information.",
        78: "The system is presented as a learning project, not a production courier platform. Features such as payment integration, live GPS devices, courier partner APIs, and enterprise deployment were kept outside the scope so that the core full stack and microservice concepts could be studied properly.",
        80: "I analyzed that courier operations need reliable user identity, accurate delivery creation, visible tracking updates, and controlled administrative access. Without a proper system, status communication can become delayed, inconsistent, or dependent on manual follow-up.",
        96: "The scope of SmartCourier was limited to a training-level delivery management system covering authentication, delivery booking, tracking events, and administrative views. I worked within these boundaries to understand how a complete application is divided into modules.",
        97: "The scope did not include real payment collection, production deployment, physical barcode scanning, live GPS integration, SMS billing, or third-party courier contracts. These exclusions kept the project realistic for an internship training assignment.",
        98: "Within the approved scope, I practiced the complete flow from frontend form submission to gateway routing, backend validation, database persistence, token-based access, and status visibility. This made the project broad enough for learning while still manageable.",
        100: "The business value of a courier management system is that it organizes delivery data, improves tracking visibility, reduces manual coordination, and helps administrators monitor work from one place.",
        101: "For my learning, the project was important because it brought together React, Spring Boot, MySQL, JWT, RabbitMQ, Eureka, API Gateway, and testing tools. Instead of studying each technology separately, I saw how they interact in one application.",
        102: "The project also taught me to think from the user's point of view. A customer needs a simple booking and tracking flow, while an administrator needs reliable summaries and controlled access to operational records.",
        104: "A SmartCourier-like system can be applied in courier companies, e-commerce delivery teams, warehouse dispatch operations, campus parcel counters, internal document movement, and small logistics service providers.",
        105: "In a real implementation, the same idea could be extended with payment gateways, label printing, courier partner APIs, GPS devices, customer notification services, proof-of-delivery capture, and analytics dashboards.",
        106: "For academic evaluation, the project is applicable because it demonstrates the essential structure of a logistics platform without requiring external hardware or commercial integrations.",
        108: "The Auth Service handles registration, login, user records, roles, password handling, and JWT generation. This module helped me understand how identity is established before a user can access protected features.",
        109: "The Delivery Service handles delivery booking, parcel details, pickup and destination information, hub references, status changes, and delivery summaries. It represents the main operational record of the system.",
        110: "The Tracking Service stores status history as separate tracking events. It helped me understand why tracking information should be maintained as a timeline rather than as only one current status field.",
        111: "The Admin Service supports dashboard and reporting functions by collecting information needed for administrative review. It can use service-to-service communication, such as OpenFeign, when it needs data owned by another module.",
        113: "By completing SmartCourier, I expected to become more confident in connecting a React frontend with secured Spring Boot APIs. I also expected to understand how module boundaries, database tables, and API contracts affect the final user flow.",
        114: "The achieved outcome was a working academic implementation that demonstrates authentication, delivery booking, tracking updates, admin views, testing, and documentation. It is suitable for explaining full stack and microservice concepts in an internship viva.",
        115: "The project improved my ability to explain architecture, test APIs, debug integration errors, write clearer documentation, and discuss practical trade-offs such as synchronous API calls versus asynchronous messaging.",

        # Chapter 3
        118: "The SmartCourier architecture combines a React single page frontend with multiple Spring Boot backend services. I studied the system by separating the user interface, gateway, service discovery, business services, messaging, and databases into clear responsibilities.",
        119: "The backend follows a microservices style in which Auth, Delivery, Tracking, and Admin responsibilities are handled by separate services. API Gateway provides the external entry point, while Eureka supports discovery between running services.",
        120: "RabbitMQ is used for event-based communication, especially for delivery status updates that can be consumed by the tracking side. This helped me understand how asynchronous processing avoids making every operation depend on one long request chain.",
        121: "JWT authentication protects secured requests. The frontend receives a token after login and sends it in the Authorization header, while the gateway and backend security configuration ensure that protected APIs are not accessed anonymously.",
        122: "The database-per-service approach keeps data ownership clear. User data, delivery records, tracking events, and administrative information are treated as separate concerns, which supports loose coupling and better maintainability.",
        134: "The React frontend is structured as a single page application with pages and feature folders for landing, authentication, customer workflows, tracking, and administration. Components are used to keep each screen focused on a specific user task.",
        135: "React Router manages navigation between public routes, customer-only routes, and admin-only routes. I used this to understand how route guards and redirects create a cleaner user experience after login.",
        136: "Context API stores authentication state such as the logged-in user and role, while Axios centralizes API calls and attaches JWT tokens to secured requests. A practical observation was that frontend checks improve usability, but final security must still be enforced by the backend.",
        138: "The Spring Boot services follow a layered structure with controllers, services, repositories, DTOs, entities, validation logic, and exception handlers. This organization made the backend easier for me to read and explain.",
        139: "Controllers handle HTTP requests and responses, service classes apply business rules, repositories access data through Spring Data JPA, and DTOs prevent direct exposure of entity objects. This separation reduced confusion when tracing a request.",
        140: "I also learned the role of annotations, dependency injection, configuration files, starter dependencies, and profiles in Spring Boot. These features helped me build APIs faster while still following a consistent project structure.",
        142: "The project applies microservices concepts by dividing authentication, delivery, tracking, and administration into services with focused responsibilities. Each service can be understood and modified without treating the whole backend as one large application.",
        143: "Independent deployment and scalability are important benefits. For example, in a real courier platform, tracking requests may be heavier than admin requests, so the tracking service could be scaled separately.",
        144: "The architecture also introduced challenges such as distributed debugging, service startup order, configuration consistency, and communication failures. These challenges helped me see why microservices require discipline, not only separate code folders.",
        146: "The API Gateway is the centralized entry point for external requests. It maps incoming paths to the correct backend service and hides individual service ports from the frontend.",
        147: "In SmartCourier, the gateway is responsible for routing, CORS handling, token checks, and request forwarding. This gave me a practical example of how cross-cutting concerns can be handled before a request reaches business logic.",
        148: "A small route mismatch can break the whole flow, so I learned to verify gateway paths, service names, Eureka registration, and JWT filter behavior together instead of checking them in isolation.",
        150: "Eureka Server provides service discovery. Services register with Eureka, and the gateway or another service can locate them using logical service names instead of fixed host and port combinations.",
        151: "This was useful for understanding dynamic backend environments. If a service port or instance changes, discovery reduces the need to hard-code connection details across the application.",
        152: "During startup, the Eureka dashboard served as a simple verification point. When a service appeared as UP, it confirmed that the application had started and registered correctly.",
        154: "RabbitMQ provides messaging support for asynchronous communication. In SmartCourier, a delivery status change can be published as an event and then processed by the Tracking Service to create tracking history.",
        155: "This pattern is useful because the Delivery Service does not need to wait for every downstream tracking operation before returning a response. It can publish the message and allow the consumer to handle tracking work separately.",
        156: "I learned the meaning of queues, exchanges, routing keys, consumers, acknowledgements, and retry concerns. These concepts showed me that messaging improves flexibility but also requires careful configuration.",
        158: "JWT authentication begins when the user logs in successfully through the Auth Service. The service generates a signed token containing identity and role information and sends it back to the frontend.",
        159: "The frontend includes the token in the Authorization header when calling protected APIs. The gateway validates the token and allows only authorized requests to proceed to customer or admin functions.",
        160: "A practical learning point was that JWT is stateless, but it still needs secure handling. Token expiry, secret management, role claims, and frontend storage decisions must be considered carefully.",
        162: "OpenFeign supports synchronous service-to-service calls through declarative Java interfaces. Instead of writing manual REST client code, a service can define the target endpoint as an interface method.",
        163: "In SmartCourier, the Admin Service can use Feign-style calls to collect delivery or user information needed for dashboard summaries. This helped me compare direct API calls with message-based communication.",
        164: "I also learned that Feign should be used carefully because one service can become dependent on another service's availability. In larger systems, fallback handling, timeouts, and circuit breakers become important.",
        166: "The database-per-service pattern means each service owns its own schema and data model. Auth owns user-related data, Delivery owns shipment records, Tracking owns event history, and Admin owns reporting-related data.",
        167: "This pattern prevents other services from directly depending on a table structure that they do not own. Changes in one module's schema can be managed within that module's boundary.",
        168: "The trade-off is that cross-service reporting cannot rely on simple database joins. Data must be obtained through APIs, events, or reporting views, which is why communication design becomes important.",
        170: "A secured request starts when the React application sends an API call to the gateway. If the user is logged in, Axios adds the JWT token to the Authorization header.",
        171: "The gateway checks the route, validates the token, and forwards the request to the appropriate service discovered through its configuration and service registry.",
        172: "The target service processes the request through controller, service, repository, and database layers. The response then returns through the gateway to the frontend, where the UI updates the relevant screen.",
        174: "The business workflow begins with registration or login. After authentication, a customer books a delivery by entering pickup, destination, receiver, and parcel information through the frontend.",
        175: "The Delivery Service stores the booking and manages status changes. When a status update occurs, an event can be sent through RabbitMQ so that the Tracking Service records the change as part of the shipment timeline.",
        176: "Administrators use their role-based access to review deliveries, users, hubs, and dashboard information. This end-to-end workflow shows how separate technical modules support one courier operation.",
        178: "The use case diagram presents the system from the viewpoint of external actors. The customer actor is linked with registration, login, delivery booking, viewing deliveries, and shipment tracking.",
        179: "The administrator actor is linked with managing users, hubs, deliveries, dashboard summaries, and reports. The diagram helps separate normal customer actions from operational control activities.",
        180: "I used the diagram during viva preparation because it gives a quick functional overview before discussing implementation details such as APIs, services, and databases.",

        # Chapter 4
        183: "For implementation, I worked with the project as an end-to-end training application rather than as isolated examples. I connected React screens with backend routes, studied how data moved through services, and documented practical behavior for viva explanation.",
        184: "On the frontend, I used React components for login, signup, customer dashboard, delivery booking, tracking, and admin screens. The main implementation concern was keeping each page understandable and ensuring navigation matched the user's role.",
        185: "For API integration, I practiced using Axios to call gateway endpoints, send JSON payloads, attach JWT headers, and handle success or error responses. This helped me understand how frontend code depends on backend contracts.",
        186: "On the backend, I studied Spring Boot controllers that expose REST APIs for authentication, delivery, tracking, and administration. Request validation, response DTOs, and exception handling were important for making the APIs predictable.",
        187: "Repositories and entities were used for persistence through Spring Data JPA and MySQL. I paid attention to fields such as user role, delivery status, tracking number, hub information, and timestamps because these fields directly affect the workflow.",
        188: "Exception handling through controller advice helped keep error responses consistent. For example, missing records, validation failures, duplicate users, and unauthorized access should produce clear responses instead of unclear backend errors.",
        192: "I worked with React components that represent practical screens. Login and Signup collect credentials, Customer Dashboard summarizes user activity, DeliveryWizard captures shipment information, and Tracking displays delivery progress.",
        193: "AdminDeliveries, UserManagement, HubManagement, and Admin Dashboard support administrator workflows. I learned that a component should not try to manage every task; smaller responsibilities make the UI easier to test and update.",
        194: "A useful example was the Tracking component, which refreshes tracking information and presents status history. This showed me how route data, API responses, loading states, and error messages combine in a real screen.",
        196: "Routing is implemented with React Router through public, customer, and admin routes. Public pages are accessible without login, while customer and admin pages are wrapped with route protection.",
        197: "The ProtectedRoute logic checks authentication status and required role. If the user is not logged in, the application redirects to login; if the role does not match, access is blocked or redirected.",
        198: "This implementation helped me understand that routing is both a user experience feature and a support mechanism for role separation. However, it must be backed by server-side checks because frontend code can be bypassed.",
        200: "Axios is used as the frontend HTTP client. I practiced creating requests for login, signup, delivery creation, delivery lists, tracking details, and dashboard data through configured backend URLs.",
        201: "After login, the token is stored and attached to protected requests using an interceptor-style approach. This made the frontend simpler because each secured call did not need to manually build the Authorization header.",
        202: "During integration, I observed that error handling matters as much as successful API calls. Invalid credentials, missing tokens, server downtime, and validation errors need clear UI handling so the user is not left confused.",
        204: "Controllers expose REST endpoints for the modules. In the Auth module, endpoints support signup, login, token validation, and user lookup. In the Delivery module, endpoints support booking, updates, listing, and status management.",
        205: "Tracking endpoints support event creation, history retrieval, latest status lookup, and filtering. Admin endpoints support dashboard and management views. These examples helped me connect HTTP methods with actual business actions.",
        206: "A practical lesson was to keep controllers thin. When request-handling code contains business rules, debugging becomes harder, so I tried to understand how controllers delegate to services.",
        208: "The service layer contains business logic such as checking duplicate users, validating delivery inputs, generating summaries, applying status changes, and preparing DTOs for responses.",
        209: "Repositories extend Spring Data JPA and provide database operations such as saving records, finding by id, counting, searching by status, and retrieving history by delivery id. This reduced boilerplate while keeping database access organized.",
        210: "The separation between services and repositories made the backend easier to test. I could reason about business decisions in the service layer and data retrieval in the repository layer separately.",
        212: "Security is implemented through JWT-based authentication and role-aware access control. I studied how the Auth Service issues tokens and how gateway validation protects routes before they reach backend services.",
        213: "The frontend uses the logged-in role to show the correct dashboard, but the server side remains responsible for real authorization. This distinction was important because hiding a button is not the same as securing an API.",
        214: "This section connected several implementation pieces: login form submission, password validation, token generation, token storage, request headers, gateway filters, and role-based access to customer or admin endpoints.",
        216: "MySQL is used as the relational database, while JPA/Hibernate maps Java entity classes to tables. I practiced how entities, repositories, and application configuration work together to persist data.",
        217: "The Auth database stores users and roles, the Delivery database stores shipment and hub information, the Tracking database stores event history, and administrative data supports dashboard functions. This reinforced the database-per-service idea.",
        218: "I also learned the value of migration scripts and repeatable setup. When a service has a clear database structure, it becomes easier to recreate the environment and explain the schema during evaluation.",

        # Chapter 5
        221: "The work completed during training was learning-oriented and practical. I used SmartCourier to connect classroom-style concepts with implementation tasks, testing, debugging, documentation, and explanation.",
        222: "My frontend work focused on React components, state handling, routing, forms, protected pages, local storage, and API integration. I learned to trace how a user action on the page becomes a request to the backend.",
        223: "My backend work focused on Spring Boot APIs, controllers, services, repositories, DTOs, validation, security, and database connectivity. I also studied how backend modules communicate in a microservice-style system.",
        224: "My microservices learning included service discovery through Eureka, routing through API Gateway, asynchronous updates through RabbitMQ, and service-to-service calls through OpenFeign. These topics helped me understand distributed application design.",
        225: "Testing and debugging became a regular part of my work. I used Postman for API testing, checked response codes and payloads, reviewed logs, and corrected configuration or integration issues step by step.",
        226: "I also participated in discussions and knowledge-sharing activities. Explaining a problem to others made me more careful about evidence, such as the exact endpoint tested, the token used, the response received, and the relevant log message.",
        232: "During training, I followed a Java Full Stack path that covered frontend development, backend services, databases, authentication, API testing, and documentation. SmartCourier acted as the common project through which these topics were connected.",
        233: "My activities included studying concepts, implementing assigned features, reviewing code behavior, preparing notes, and correcting mistakes found during testing. The main outcome was not speed, but a better understanding of how full stack layers depend on one another.",
        235: "In React, I practiced building components, using hooks, handling form state, applying Context API for authentication state, and configuring routes for customer and admin users.",
        236: "I learned that a screen should guide the user clearly. For example, login should show meaningful errors, booking should capture required fields in order, and tracking should present status history in a readable way.",
        238: "In Spring Boot, I learned to create REST APIs using annotations, controllers, services, repositories, DTOs, validation classes, and exception handlers. I also understood how dependency injection keeps classes easier to manage.",
        239: "A major learning point was keeping business logic in the service layer. This made the code easier to explain because controllers handled requests, services handled decisions, and repositories handled persistence.",
        241: "I practiced API development by working with endpoints for authentication, delivery, tracking, and administration. I focused on request methods, path variables, request bodies, response DTOs, and suitable status codes.",
        242: "I tested normal cases first, then moved to invalid data, missing tokens, wrong roles, unknown ids, and incorrect payloads. This helped me understand that a good API must handle errors clearly, not only successful requests.",
        244: "I gained practical exposure to database connectivity through Spring Boot configuration, entity classes, repositories, and migration scripts. I saw how small configuration issues can stop a service from starting or saving data.",
        245: "A key observation was that database design affects the whole application. Fields such as role, tracking number, delivery status, hub id, and timestamps must be defined carefully because they appear in validation, queries, and UI displays.",
        247: "I studied authentication through signup, login, password handling, JWT generation, token storage, and gateway validation. This helped me understand how identity travels from the login page to protected backend APIs.",
        248: "I also practiced role-based behavior. Customer and admin users require different screens and different API access, so both frontend route checks and backend authorization checks are needed.",
        249: "Testing Activities",
        250: "I used Postman to test endpoints manually and observe request-response behavior. My testing included authorization headers, JSON bodies, expected status codes, invalid inputs, and role-based access checks.",
        251: "Testing helped me isolate issues. If an API worked in Postman but not in React, I checked frontend request URLs or headers; if it failed in Postman too, I checked backend logs, routing, validation, or database state.",
        252: "Debugging Experience",
        253: "My debugging experience included checking application logs, reviewing YAML configuration, confirming service ports, verifying Eureka registration, and testing gateway routes one by one.",
        254: "Common issues involved missing bearer tokens, route prefix mismatches, database connection settings, validation failures, RabbitMQ startup, or service startup order. These tasks improved my patience and my habit of testing one assumption at a time.",
        256: "I participated in discussions on module behavior, API design, test scenarios, and project explanation. These interactions showed me that software work also depends on communication and shared understanding.",
        257: "I learned to describe problems more clearly by stating what I tested, what result I expected, what actually happened, and which response or log supported my conclusion.",
        259: "Knowledge-sharing sessions helped me revise concepts such as REST APIs, Spring Boot layers, JWT, RabbitMQ, Eureka, React routing, and database-per-service design. I used these sessions to connect theory with the SmartCourier implementation.",
        260: "They also improved my viva preparation because I had to explain the same concept in simple language. For example, I practiced explaining why RabbitMQ is asynchronous and why Eureka avoids hard-coded service addresses.",
        262: "The main challenges I faced were integration-related. A feature could appear correct in the frontend but fail because of gateway routing, token headers, backend validation, or a service not being registered.",
        263: "Another challenge was understanding how microservices communicate without sharing one database. I had to think carefully about which service owns which data and how another service should request or receive it.",
        265: "I solved issues by checking each layer separately: frontend request, gateway route, token validation, controller endpoint, service logic, repository query, and database record. This step-by-step method reduced guesswork.",
        266: "For repeated issues, I prepared notes and test cases so that I could reproduce the problem and verify the fix. This habit was useful while preparing the final report and viva explanations.",
        268: "By the end of training, I improved my understanding of React, Spring Boot, REST APIs, MySQL, JWT, RabbitMQ, Eureka, OpenFeign, testing, debugging, and documentation.",
        269: "More importantly, I learned how to explain my work honestly as a trainee: what I implemented, what I studied, what limitations remained, and how the project could be improved in a real production environment.",

        # Lightly revise later chapters to reduce repeated rhythm
        272: "Testing was performed to confirm whether the main SmartCourier workflows behaved as expected. I treated testing as part of development rather than as a final activity after coding.",
        273: "The testing work covered login, signup, protected route access, delivery creation, delivery listing, tracking history, admin access, validation failures, and selected service-level behavior.",
        274: "Postman was used for API verification, while unit and controller tests in the backend helped confirm service behavior for important cases. Frontend tests supported selected UI interactions.",
        275: "The testing process also helped me find documentation gaps. When an endpoint or expected response was unclear, I updated my notes so that the workflow could be explained more accurately.",
        300: "The SmartCourier project gave me practical exposure to full stack development in a microservice-oriented training environment. It helped me move from understanding individual technologies to explaining how they work together.",
        301: "The project covered user authentication, delivery booking, tracking history, administrative views, API routing, service discovery, messaging, database connectivity, testing, and documentation.",
        302: "My contribution was realistic for an intern: I studied the architecture, implemented and tested training-level features, debugged integration issues, prepared documentation, and learned to explain the system clearly.",
        303: "The final outcome is an academic project that demonstrates the main concepts of a courier management platform while acknowledging that production deployment would require more security, scalability, monitoring, and external integrations.",
    }

    for idx, text in replacements.items():
        if idx in (249, 252):
            replace_heading(doc.paragraphs[idx], text)
        else:
            replace_paragraph(doc.paragraphs[idx], text)

    # Rename headings to match the requested section labels while preserving styles.
    rename = {
        249: "Testing Activities",
        252: "Debugging Experience",
    }
    for idx, text in rename.items():
        replace_heading(doc.paragraphs[idx], text)

    # Rewrite the weekly table, which had repeated placeholder wording.
    weekly_rows = [
        ("Week 1", "Java refresh and OOP", "I revised classes, objects, inheritance, interfaces, collections, and exception handling through short coding exercises."),
        ("Week 2", "Spring Boot basics", "I created simple REST endpoints and learned how annotations, dependency injection, and application configuration work."),
        ("Week 3", "REST API design", "I practiced request methods, path variables, request bodies, response DTOs, and status code handling."),
        ("Week 4", "Database and JPA", "I connected services with MySQL, mapped entities, used repositories, and checked how database records changed after API calls."),
        ("Week 5", "React fundamentals", "I built reusable components, handled state with hooks, and improved form handling for login and delivery screens."),
        ("Week 6", "Routing and auth state", "I worked with React Router, protected routes, Context API, local storage, and role-based navigation."),
        ("Week 7", "JWT authentication", "I studied signup, login, token generation, Authorization headers, gateway validation, and role-based access checks."),
        ("Week 8", "Delivery module", "I focused on delivery booking, status changes, delivery lists, hub references, and response mapping."),
        ("Week 9", "Tracking module", "I worked with tracking events, latest status retrieval, timeline display, and event history for each delivery."),
        ("Week 10", "API Gateway", "I verified route prefixes, service forwarding, CORS behavior, and the impact of gateway configuration errors."),
        ("Week 11", "Eureka discovery", "I observed service registration, checked service availability, and understood why logical service names are useful."),
        ("Week 12", "RabbitMQ messaging", "I studied exchanges, queues, routing keys, consumers, and asynchronous status update handling."),
        ("Week 13", "OpenFeign and admin flow", "I learned how a service can request data from another service for dashboard or summary information."),
        ("Week 14", "Testing practice", "I used Postman and backend tests to verify success cases, validation failures, missing tokens, and role restrictions."),
        ("Week 15", "Debugging practice", "I traced integration issues across frontend requests, gateway routes, service logs, database settings, and message flow."),
        ("Week 16", "Documentation", "I prepared notes, diagrams, explanations, and report content in a way suitable for internship evaluation."),
        ("Week 17", "Viva preparation", "I practiced explaining architecture, modules, implementation choices, limitations, and future enhancements."),
        ("Week 18", "Review and refinement", "I reviewed the project end to end, corrected wording in documentation, and organized the final deliverables."),
        ("Week 19", "Final consolidation", "I connected the learning outcomes with the completed SmartCourier workflows and prepared for submission."),
        ("Week 20", "Submission readiness", "I checked the report, presentation, testing notes, and viva answers for consistency and clarity."),
    ]
    table = doc.tables[12]
    for row_idx, row_data in enumerate(weekly_rows, start=1):
        for col_idx, value in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = value

    # Improve training activity table where rows were too generic.
    activity_rows = [
        ("Orientation", "I reviewed training objectives, tools, expectations, and the SmartCourier project idea.", "I understood the purpose of the internship work and prepared my setup."),
        ("React Practice", "I built screens, forms, routes, protected views, and basic dashboard flows.", "I improved my understanding of component-driven frontend development."),
        ("Spring Boot Practice", "I worked with controllers, services, repositories, DTOs, validation, and exception handling.", "I learned how backend code is organized in maintainable layers."),
        ("Database Work", "I connected services to MySQL and observed how entity fields are stored and queried.", "I understood the relationship between API behavior and database design."),
        ("Security Work", "I studied signup, login, JWT generation, token headers, and role-based access.", "I learned why frontend route protection must be supported by backend validation."),
        ("Testing and Debugging", "I tested APIs in Postman, checked logs, and isolated issues across gateway, services, and database layers.", "I developed a more systematic approach to finding and explaining errors."),
        ("Documentation", "I prepared report sections, diagrams, module explanations, and viva notes.", "I improved my ability to present technical work in an academic format."),
    ]
    table = doc.tables[11]
    for row_idx, row_data in enumerate(activity_rows, start=1):
        for col_idx, value in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = value

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
