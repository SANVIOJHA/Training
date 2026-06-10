from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


DOCX = Path(r"D:\Viva_etp\SmartCourier_LPU_Training_Report_.docx")
OUTPUT_DOCX = Path(r"D:\Viva_etp\SmartCourier_LPU_Training_Report_With_Certificate.docx")
IMAGE = Path(r"C:\Users\ASUS\OneDrive\Pictures\ScreenShotss\Screenshot 2026-06-03 223609.png")


def has_page_break(paragraph) -> bool:
    return any(run._element.xpath('.//w:br[@w:type="page"]') for run in paragraph.runs)


def insert_paragraph_before(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def insert_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def main():
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)
    if not IMAGE.exists():
        raise FileNotFoundError(IMAGE)

    doc = Document(str(DOCX))
    cert_index = next(
        i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Training Certificate"
    )

    # Remove filler blank paragraphs immediately before the certificate heading.
    for p in list(doc.paragraphs[:cert_index])[::-1]:
        if p.text.strip() or has_page_break(p):
            break
        remove_paragraph(p)

    cert_index = next(
        i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Training Certificate"
    )
    cert = doc.paragraphs[cert_index]
    previous = doc.paragraphs[cert_index - 1]
    if not has_page_break(previous):
        break_p = insert_paragraph_before(cert)
        break_p.add_run().add_break(WD_BREAK.PAGE)

    cert_index = next(
        i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Training Certificate"
    )
    cert = doc.paragraphs[cert_index]
    next_idx = cert_index + 1
    if next_idx < len(doc.paragraphs):
        next_para = doc.paragraphs[next_idx]
        if not has_page_break(next_para):
            remove_paragraph(next_para)

    image_p = insert_paragraph_after(cert)
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.paragraph_format.space_before = Pt(4)
    image_p.paragraph_format.space_after = Pt(0)
    image_p.add_run().add_picture(str(IMAGE), width=Inches(5.65))

    doc.save(str(OUTPUT_DOCX))
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
